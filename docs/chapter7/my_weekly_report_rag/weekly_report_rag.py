#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""周报 RAG 系统：单文件实现，支持 docx 解析、向量检索与问答。"""

import argparse
import json
import os
import re
import shutil
import sys
from typing import List

import numpy as np
import tiktoken
from docx import Document
from dotenv import find_dotenv, load_dotenv
from openai import OpenAI
from tqdm import tqdm

_ = load_dotenv(find_dotenv())

# ── 配置 ──────────────────────────────────────────────────────────────────────
REPORT_DATA_PATH = "/Users/bigo/Desktop/bigo/bigo工作周报"
STORAGE_PATH = "./storage"
EMBEDDING_MODEL = "BAAI/bge-m3"
CHAT_MODEL = "Qwen/Qwen2.5-32B-Instruct"
MAX_TOKEN_LEN = 600
COVER_CONTENT = 150
DEFAULT_K = 3

WEEKLY_REPORT_PROMPT = """
你是工作周报助手。根据以下周报片段回答用户问题。
要求：用中文回答；引用具体日期/项目/数据；上下文不足时说"周报中没有相关内容"。

问题: {question}
可参考的周报内容：
···
{context}
···
回答:
"""

enc = tiktoken.get_encoding("cl100k_base")


def _get_token_cover(text: str, cover_tokens: int) -> str:
    tokens = enc.encode(text)
    if len(tokens) <= cover_tokens:
        return text
    return enc.decode(tokens[-cover_tokens:])


# ── 文档加载与切分 ────────────────────────────────────────────────────────────

class ReadFiles:
    def __init__(self, path: str) -> None:
        self._path = os.path.abspath(path)
        self.file_list = self.get_files()

    def get_files(self) -> List[str]:
        file_list = []
        for filepath, _, filenames in os.walk(self._path):
            for filename in filenames:
                if filename.startswith("~$") or filename == "tmp.docx":
                    continue
                if filename.endswith(".docx"):
                    file_list.append(os.path.join(filepath, filename))
        return sorted(file_list)

    def get_content(self, max_token_len: int = MAX_TOKEN_LEN, cover_content: int = COVER_CONTENT) -> List[str]:
        docs = []
        for file in self.file_list:
            content = self.read_file_content(file)
            if not content.strip():
                continue
            rel_path = os.path.relpath(file, self._path)
            chunks = self.get_chunk(content, max_token_len=max_token_len, cover_content=cover_content)
            for chunk in chunks:
                docs.append(f"[来源: {rel_path}]\n{chunk}")
        return docs

    @classmethod
    def get_chunk(cls, text: str, max_token_len: int = MAX_TOKEN_LEN, cover_content: int = COVER_CONTENT) -> List[str]:
        chunk_text = []
        curr_len = 0
        curr_chunk = ""
        token_len = max_token_len - cover_content
        lines = text.splitlines()

        for line in lines:
            line = line.strip()
            if not line:
                continue
            line_len = len(enc.encode(line))

            if line_len > max_token_len:
                if curr_chunk:
                    chunk_text.append(curr_chunk)
                    curr_chunk = ""
                    curr_len = 0

                line_tokens = enc.encode(line)
                num_chunks = (len(line_tokens) + token_len - 1) // token_len

                for i in range(num_chunks):
                    start_token = i * token_len
                    end_token = min(start_token + token_len, len(line_tokens))
                    chunk_part = enc.decode(line_tokens[start_token:end_token])

                    if i > 0:
                        overlap_start = max(0, start_token - cover_content)
                        cover_part = enc.decode(line_tokens[overlap_start:start_token])
                        chunk_part = cover_part + chunk_part

                    chunk_text.append(chunk_part)

                curr_chunk = ""
                curr_len = 0

            elif curr_len + line_len + 1 <= token_len:
                if curr_chunk:
                    curr_chunk += "\n"
                    curr_len += 1
                curr_chunk += line
                curr_len += line_len
            else:
                if curr_chunk:
                    chunk_text.append(curr_chunk)
                if chunk_text:
                    cover_part = _get_token_cover(chunk_text[-1], cover_content)
                    curr_chunk = cover_part + "\n" + line
                    curr_len = len(enc.encode(cover_part)) + 1 + line_len
                else:
                    curr_chunk = line
                    curr_len = line_len

        if curr_chunk:
            chunk_text.append(curr_chunk)

        return chunk_text

    @classmethod
    def read_file_content(cls, file_path: str) -> str:
        if file_path.endswith(".docx"):
            return cls.read_docx(file_path)
        raise ValueError(f"Unsupported file type: {file_path}")

    @classmethod
    def read_docx(cls, file_path: str) -> str:
        doc = Document(file_path)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)


# ── Embedding ─────────────────────────────────────────────────────────────────

class BaseEmbeddings:
    def __init__(self, path: str = "", is_api: bool = True) -> None:
        self.path = path
        self.is_api = is_api

    def get_embedding(self, text: str, model: str) -> List[float]:
        raise NotImplementedError

    @classmethod
    def cosine_similarity(cls, vector1: List[float], vector2: List[float]) -> float:
        v1 = np.array(vector1, dtype=np.float32)
        v2 = np.array(vector2, dtype=np.float32)
        if not np.all(np.isfinite(v1)) or not np.all(np.isfinite(v2)):
            return 0.0
        dot_product = np.dot(v1, v2)
        magnitude = np.linalg.norm(v1) * np.linalg.norm(v2)
        if magnitude == 0:
            return 0.0
        return dot_product / magnitude


class OpenAIEmbedding(BaseEmbeddings):
    def __init__(self, path: str = "", is_api: bool = True) -> None:
        super().__init__(path, is_api)
        if self.is_api:
            self.client = OpenAI()
            self.client.api_key = os.getenv("OPENAI_API_KEY")
            self.client.base_url = os.getenv("OPENAI_BASE_URL")

    def get_embedding(self, text: str, model: str = EMBEDDING_MODEL) -> List[float]:
        if self.is_api:
            text = text.replace("\n", " ")
            return self.client.embeddings.create(input=[text], model=model).data[0].embedding
        raise NotImplementedError


# ── 向量存储 ──────────────────────────────────────────────────────────────────

class VectorStore:
    def __init__(self, document: List[str] = None) -> None:
        self.document = document or []
        self.vectors: List[List[float]] = []

    def get_vector(self, embedding_model: BaseEmbeddings) -> List[List[float]]:
        self.vectors = []
        for doc in tqdm(self.document, desc="Calculating embeddings"):
            self.vectors.append(embedding_model.get_embedding(doc))
        return self.vectors

    def persist(self, path: str = STORAGE_PATH) -> None:
        os.makedirs(path, exist_ok=True)
        with open(f"{path}/document.json", "w", encoding="utf-8") as f:
            json.dump(self.document, f, ensure_ascii=False)
        if self.vectors:
            with open(f"{path}/vectors.json", "w", encoding="utf-8") as f:
                json.dump(self.vectors, f)

    def load_vector(self, path: str = STORAGE_PATH) -> None:
        with open(f"{path}/vectors.json", "r", encoding="utf-8") as f:
            self.vectors = json.load(f)
        with open(f"{path}/document.json", "r", encoding="utf-8") as f:
            self.document = json.load(f)

    def get_similarity(self, vector1: List[float], vector2: List[float]) -> float:
        return BaseEmbeddings.cosine_similarity(vector1, vector2)

    def query(self, query: str, embedding_model: BaseEmbeddings, k: int = DEFAULT_K) -> List[str]:
        query_vector = embedding_model.get_embedding(query)
        result = np.array([self.get_similarity(query_vector, vector) for vector in self.vectors])
        return np.array(self.document)[result.argsort()[-k:][::-1]].tolist()


# ── LLM 问答 ─────────────────────────────────────────────────────────────────

class OpenAIChat:
    def __init__(self, model: str = CHAT_MODEL) -> None:
        self.model = model

    def chat(self, prompt: str, history: List[dict], content: str) -> str:
        client = OpenAI()
        client.api_key = os.getenv("OPENAI_API_KEY")
        client.base_url = os.getenv("OPENAI_BASE_URL")
        history.append({
            "role": "user",
            "content": WEEKLY_REPORT_PROMPT.format(question=prompt, context=content),
        })
        response = client.chat.completions.create(
            model=self.model,
            messages=history,
            max_tokens=2048,
            temperature=0.1,
        )
        return response.choices[0].message.content


# ── 业务逻辑 ──────────────────────────────────────────────────────────────────

def _check_env() -> None:
    if not os.getenv("OPENAI_API_KEY") or not os.getenv("OPENAI_BASE_URL"):
        print("错误: 请在 .env 中配置 OPENAI_API_KEY 和 OPENAI_BASE_URL")
        sys.exit(1)


def build_index(data_path: str, storage_path: str, force: bool = False) -> VectorStore:
    _check_env()
    if os.path.exists(storage_path) and not force:
        print(f"storage 已存在: {storage_path}，使用 --force 强制重建")
        sys.exit(1)

    if force and os.path.exists(storage_path):
        shutil.rmtree(storage_path)

    reader = ReadFiles(data_path)
    print(f"扫描到 {len(reader.file_list)} 个 docx 文件")
    if not reader.file_list:
        print(f"错误: 在 {data_path} 下未找到 docx 文件")
        sys.exit(1)

    docs = reader.get_content()
    print(f"切分为 {len(docs)} 个 chunk")

    vector = VectorStore(docs)
    embedding = OpenAIEmbedding()
    vector.get_vector(embedding_model=embedding)
    vector.persist(path=storage_path)
    print(f"向量库已保存到 {storage_path}")
    return vector


def load_index(storage_path: str) -> VectorStore:
    vectors_file = os.path.join(storage_path, "vectors.json")
    document_file = os.path.join(storage_path, "document.json")
    if not os.path.exists(vectors_file) or not os.path.exists(document_file):
        print(f"错误: storage 不存在或不完整，请先运行: python weekly_report_rag.py --build")
        sys.exit(1)
    vector = VectorStore()
    vector.load_vector(storage_path)
    return vector


def ask(question: str, storage_path: str, k: int = DEFAULT_K) -> str:
    _check_env()
    vector = load_index(storage_path)
    embedding = OpenAIEmbedding()
    contexts = vector.query(question, embedding_model=embedding, k=k)
    context = "\n\n---\n\n".join(contexts)
    chat = OpenAIChat()
    return chat.chat(question, [], context)


def interactive_chat(storage_path: str, k: int = DEFAULT_K) -> None:
    _check_env()
    vector = load_index(storage_path)
    embedding = OpenAIEmbedding()
    chat = OpenAIChat()
    history: List[dict] = []

    print("周报 RAG 交互模式（输入 quit 退出）")
    while True:
        try:
            question = input("\n问题> ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n再见")
            break
        if not question:
            continue
        if question.lower() in ("quit", "exit", "q"):
            print("再见")
            break

        contexts = vector.query(question, embedding_model=embedding, k=k)
        context = "\n\n---\n\n".join(contexts)
        answer = chat.chat(question, history.copy(), context)
        print(f"\n{answer}")
        history.append({"role": "user", "content": question})
        history.append({"role": "assistant", "content": answer})


def main() -> None:
    parser = argparse.ArgumentParser(description="周报 RAG 系统")
    parser.add_argument("--build", action="store_true", help="构建向量索引")
    parser.add_argument("--force", action="store_true", help="强制重建索引（覆盖已有 storage）")
    parser.add_argument("--query", type=str, help="单次问答")
    parser.add_argument("--chat", action="store_true", help="交互式问答")
    parser.add_argument("--k", type=int, default=DEFAULT_K, help=f"检索 top-k 片段（默认 {DEFAULT_K}）")
    parser.add_argument("--data-path", type=str, default=REPORT_DATA_PATH, help="周报原始数据路径")
    parser.add_argument("--storage", type=str, default=STORAGE_PATH, help="向量库保存路径")
    args = parser.parse_args()

    if args.build:
        build_index(args.data_path, args.storage, force=args.force)
    elif args.query:
        print(ask(args.query, args.storage, k=args.k))
    elif args.chat:
        interactive_chat(args.storage, k=args.k)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
