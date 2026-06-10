#!/usr/bin/env python
# -*- coding: utf-8 -*-

import json
import os
import re
from typing import List, Optional, Tuple

from docx import Document

from chunker import count_tokens, get_chunk
from config import HEADING_MAX_LEN, MAX_TOKEN_LEN, PARSER_PROJECT_HINTS, PARSER_RULES_PATH, PARSER_SECTION_PREFIXES
from models import ChunkMetadata, DocumentChunk
from utils import (
    format_report_date,
    normalize_report_date,
    parse_author_from_path,
    parse_quarter_from_path,
    parse_report_date_from_path,
    parse_report_date_from_text,
)


def load_parser_rules() -> dict:
    with open(PARSER_RULES_PATH, encoding="utf-8") as handle:
        return json.load(handle)


def is_report_title(line: str, rules: dict) -> bool:
    keyword = rules.get("report_title_contains", "工作周报")
    return keyword in line and bool(re.search(r"\d{4}", line))


def is_heading_style(style_name: Optional[str], rules: dict) -> bool:
    if not style_name:
        return False
    keywords = rules.get("heading_style_keywords", [])
    return any(keyword.lower() in style_name.lower() for keyword in keywords)


def matches_project_hint(line: str) -> bool:
    lower = line.lower()
    return any(hint.lower() in lower for hint in PARSER_PROJECT_HINTS)


def is_section_heading(
    line: str,
    next_line: Optional[str],
    style_name: Optional[str] = None,
    rules: Optional[dict] = None,
) -> bool:
    rules = rules or load_parser_rules()
    line = line.strip()
    if not line or line.startswith("http"):
        return False
    if is_report_title(line, rules):
        return False
    if is_heading_style(style_name, rules):
        return True
    if rules.get("bracket_heading") and re.match(r"^【.+】$", line):
        return True
    if matches_project_hint(line) and len(line) <= HEADING_MAX_LEN:
        return True
    if len(line) > HEADING_MAX_LEN:
        return False
    endings = tuple(rules.get("sentence_endings", []))
    if line.endswith(endings) and len(line) > 15:
        return False
    if re.match(r"^\d{1,2}月\d{1,2}日", line):
        return False
    for prefixes in PARSER_SECTION_PREFIXES.values():
        if any(line.lower().startswith(prefix.lower()) for prefix in prefixes):
            return False
    if len(line) <= HEADING_MAX_LEN:
        if next_line is None:
            return len(line) <= 20
        if len(next_line) <= HEADING_MAX_LEN and not next_line.endswith("。"):
            return True
        if len(next_line) > len(line) + 10 or next_line.endswith("。"):
            return True
    return False


def infer_section_type(body_lines: List[str]) -> str:
    for line in body_lines[:5]:
        stripped = line.strip()
        for section_type, prefixes in PARSER_SECTION_PREFIXES.items():
            if any(stripped.startswith(prefix) for prefix in prefixes):
                return section_type
    return "body"


def build_chunk_text(metadata: ChunkMetadata, body_lines: List[str]) -> str:
    header = (
        f"[来源: {metadata.source} | 日期: {format_report_date(metadata.report_date)} | "
        f"项目: {metadata.project} | 季度: {metadata.quarter or '-'} | "
        f"作者: {metadata.author or '-'} | 类型: {metadata.section_type}]"
    )
    body = "\n".join([metadata.project] + body_lines) if body_lines else metadata.project
    return f"{header}\n{body}"


def split_docx_into_sections(file_path: str) -> List[Tuple[str, List[str]]]:
    rules = load_parser_rules()
    doc = Document(file_path)
    paragraphs: List[Tuple[str, Optional[str]]] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else None
        paragraphs.append((text, style_name))

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append((row_text, None))

    sections: List[Tuple[str, List[str]]] = []
    current_heading = "综合"
    current_lines: List[str] = []

    for index, (line, style_name) in enumerate(paragraphs):
        if is_report_title(line, rules):
            continue
        next_line = paragraphs[index + 1][0] if index + 1 < len(paragraphs) else None
        if is_section_heading(line, next_line, style_name, rules):
            if current_lines or current_heading != "综合":
                sections.append((current_heading, current_lines))
            current_heading = line
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_heading, current_lines))

    if not sections and paragraphs:
        sections.append(("综合", [item[0] for item in paragraphs]))

    return sections


def sections_to_chunks(file_path: str, rel_path: str, report_date: str, sections: List[Tuple[str, List[str]]]) -> List[DocumentChunk]:
    chunks: List[DocumentChunk] = []
    author = parse_author_from_path(file_path)
    quarter = parse_quarter_from_path(file_path)
    report_date = normalize_report_date(report_date)

    for project, body_lines in sections:
        if not body_lines:
            continue
        base_metadata = ChunkMetadata(
            source=rel_path,
            report_date=report_date,
            project=project,
            author=author,
            quarter=quarter,
            section_type=infer_section_type(body_lines),
        )
        section_text = build_chunk_text(base_metadata, body_lines)
        if count_tokens(section_text) <= MAX_TOKEN_LEN:
            chunks.append(DocumentChunk(section_text, base_metadata))
        else:
            for chunk_index, part in enumerate(get_chunk(section_text)):
                metadata = ChunkMetadata(
                    source=rel_path,
                    report_date=report_date,
                    project=project,
                    chunk_index=chunk_index,
                    author=author,
                    quarter=quarter,
                    section_type=base_metadata.section_type,
                )
                chunks.append(DocumentChunk(part, metadata))
    return chunks


class ReadFiles:
    def __init__(self, path: str) -> None:
        self.data_path = os.path.abspath(path)
        self.file_list = self._scan_files()

    def _scan_files(self) -> List[str]:
        file_list = []
        for filepath, _, filenames in os.walk(self.data_path):
            for filename in filenames:
                if filename.startswith("~$") or filename == "tmp.docx":
                    continue
                if filename.endswith(".docx"):
                    file_list.append(os.path.join(filepath, filename))
        return sorted(file_list)

    def _resolve_report_date(
        self,
        file_path: str,
        sections: List[Tuple[str, List[str]]],
    ) -> str:
        report_date = parse_report_date_from_path(file_path)
        if not report_date and sections:
            first_line = sections[0][1][0] if sections[0][1] else sections[0][0]
            report_date = parse_report_date_from_text(first_line)
        return report_date

    def get_chunks_for_file(self, file_path: str) -> List[DocumentChunk]:
        rel_path = os.path.relpath(file_path, self.data_path)
        sections = split_docx_into_sections(file_path)
        report_date = self._resolve_report_date(file_path, sections)
        return sections_to_chunks(file_path, rel_path, report_date, sections)

    def get_chunks(self) -> List[DocumentChunk]:
        chunks: List[DocumentChunk] = []
        for file_path in self.file_list:
            chunks.extend(self.get_chunks_for_file(file_path))
        return chunks
