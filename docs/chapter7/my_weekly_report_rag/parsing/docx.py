#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""docx 周报解析：识别项目标题、切分章节并生成 DocumentChunk。"""

import re
from typing import List, Optional, Tuple

from docx import Document

from config import MAX_TOKEN_LEN
from models import ChunkMetadata, DocumentChunk
from parsing.chunker import count_tokens, get_chunk
from parsing.rules import heading_max_len, load_parser_rules, parser_project_hints, section_prefixes
from utils import (
    format_report_date,
    normalize_report_date,
    parse_author_from_path,
    parse_quarter_from_path,
)


def is_report_title(line: str, rules: dict) -> bool:
    """判断是否为周报总标题行（非项目章节）。"""
    keyword = rules.get("report_title_contains", "工作周报")
    return keyword in line and bool(re.search(r"\d{4}", line))


def is_heading_style(style_name: Optional[str], rules: dict) -> bool:
    """根据 Word 样式名判断是否为标题。"""
    if not style_name:
        return False
    keywords = rules.get("heading_style_keywords", [])
    return any(keyword.lower() in style_name.lower() for keyword in keywords)


def matches_project_hint(line: str) -> bool:
    """行文本是否包含已知项目名关键词。"""
    lower = line.lower()
    return any(hint.lower() in lower for hint in parser_project_hints())


def is_section_heading(
    line: str,
    next_line: Optional[str],
    style_name: Optional[str] = None,
    rules: Optional[dict] = None,
) -> bool:
    """综合样式、长度、上下文启发式判断是否为项目/章节标题。"""
    rules = rules or load_parser_rules()
    max_len = heading_max_len()
    line = line.strip()
    if not line or line.startswith("http"):
        return False
    if is_report_title(line, rules):
        return False
    if is_heading_style(style_name, rules):
        return True
    if rules.get("bracket_heading") and re.match(r"^【.+】$", line):
        return True
    if matches_project_hint(line) and len(line) <= max_len:
        return True
    if len(line) > max_len:
        return False
    endings = tuple(rules.get("sentence_endings", []))
    if line.endswith(endings) and len(line) > 15:
        return False
    if re.match(r"^\d{1,2}月\d{1,2}日", line):
        return False
    for prefixes in section_prefixes().values():
        if any(line.lower().startswith(prefix.lower()) for prefix in prefixes):
            return False
    if len(line) <= max_len:
        if next_line is None:
            return len(line) <= 20
        if len(next_line) <= max_len and not next_line.endswith("。"):
            return True
        if len(next_line) > len(line) + 10 or next_line.endswith("。"):
            return True
    return False


def infer_section_type(body_lines: List[str]) -> str:
    """根据段落前缀推断章节类型（如 plan、risk、body）。"""
    prefixes_map = section_prefixes()
    for line in body_lines[:5]:
        stripped = line.strip()
        for section_type, prefixes in prefixes_map.items():
            if any(stripped.startswith(prefix) for prefix in prefixes):
                return section_type
    return "body"


def build_chunk_text(metadata: ChunkMetadata, body_lines: List[str]) -> str:
    """组装 chunk 文本：首行 meta 头 + 项目名 + 正文。"""
    header = (
        f"[来源: {metadata.source} | 日期: {format_report_date(metadata.report_date)} | "
        f"项目: {metadata.project} | 季度: {metadata.quarter or '-'} | "
        f"作者: {metadata.author or '-'} | 类型: {metadata.section_type}]"
    )
    body = "\n".join([metadata.project] + body_lines) if body_lines else metadata.project
    return f"{header}\n{body}"


def split_docx_into_sections(file_path: str) -> List[Tuple[str, List[str]]]:
    """将 docx 按项目标题切分为 (项目名, 正文行列表) 列表。"""
    rules = load_parser_rules()
    doc = Document(file_path)
    paragraphs: List[Tuple[str, Optional[str]]] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else None
        paragraphs.append((text, style_name))

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append((row_text, None))

    sections: List[Tuple[str, List[str]]] = []
    current_heading = "综合"
    current_lines: List[str] = []

    for index, (line, style_name) in enumerate(paragraphs):
        if is_report_title(line, rules):
            continue
        next_line = paragraphs[index + 1][0] if index + 1 < len(paragraphs) else None
        if is_section_heading(line, next_line, style_name, rules):
            if current_lines or current_heading != "综合":
                sections.append((current_heading, current_lines))
            current_heading = line
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_heading, current_lines))

    if not sections and paragraphs:
        sections.append(("综合", [item[0] for item in paragraphs]))

    return sections


def sections_to_chunks(
    file_path: str,
    rel_path: str,
    report_date: str,
    sections: List[Tuple[str, List[str]]],
) -> List[DocumentChunk]:
    """将章节列表转为 DocumentChunk，超长章节进一步切分。"""
    chunks: List[DocumentChunk] = []
    author = parse_author_from_path(file_path)
    quarter = parse_quarter_from_path(file_path)
    report_date = normalize_report_date(report_date)

    for project, body_lines in sections:
        if not body_lines:
            continue
        base_metadata = ChunkMetadata(
            source=rel_path,
            report_date=report_date,
            project=project,
            author=author,
            quarter=quarter,
            section_type=infer_section_type(body_lines),
        )
        section_text = build_chunk_text(base_metadata, body_lines)
        if count_tokens(section_text) <= MAX_TOKEN_LEN:
            chunks.append(DocumentChunk(section_text, base_metadata))
        else:
            for chunk_index, part in enumerate(get_chunk(section_text)):
                metadata = ChunkMetadata(
                    source=rel_path,
                    report_date=report_date,
                    project=project,
                    chunk_index=chunk_index,
                    author=author,
                    quarter=quarter,
                    section_type=base_metadata.section_type,
                )
                chunks.append(DocumentChunk(part, metadata))
    return chunks
