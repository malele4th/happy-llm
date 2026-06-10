#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import re
from dataclasses import dataclass
from typing import List, Optional, Tuple

import tiktoken
from docx import Document

from config import COVER_CONTENT, HEADING_MAX_LEN, MAX_TOKEN_LEN

enc = tiktoken.get_encoding("cl100k_base")


@dataclass
class DocumentChunk:
    text: str
    source: str
    report_date: str
    project: str

    def to_metadata(self) -> dict:
        return {
            "source": self.source,
            "report_date": self.report_date,
            "project": self.project,
        }


def get_token_cover(text: str, cover_tokens: int) -> str:
    tokens = enc.encode(text)
    if len(tokens) <= cover_tokens:
        return text
    return enc.decode(tokens[-cover_tokens:])


def parse_report_date_from_path(file_path: str) -> str:
    match = re.search(r"(\d{8})", os.path.basename(file_path))
    return match.group(1) if match else ""


def parse_report_date_from_text(text: str) -> str:
    match = re.search(r"(\d{4})-(\d{2})-(\d{2})", text)
    if match:
        return f"{match.group(1)}{match.group(2)}{match.group(3)}"
    match = re.search(r"(\d{8})", text)
    return match.group(1) if match else ""


def format_report_date(report_date: str) -> str:
    if len(report_date) == 8:
        return f"{report_date[:4]}-{report_date[4:6]}-{report_date[6:8]}"
    return report_date


def parse_date_filter(question: str) -> Tuple[Optional[int], Optional[int]]:
    match = re.search(r"(\d{4})年(\d{1,2})月", question)
    if match:
        return int(match.group(1)), int(match.group(2))
    match = re.search(r"(\d{4})年", question)
    if match:
        return int(match.group(1)), None
    match = re.search(r"(\d{8})", question)
    if match:
        d = match.group(1)
        return int(d[:4]), int(d[4:6])
    return None, None


def is_report_title(line: str) -> bool:
    return "工作周报" in line and bool(re.search(r"\d{4}", line))


def is_section_heading(line: str, next_line: Optional[str]) -> bool:
    line = line.strip()
    if not line or line.startswith("http"):
        return False
    if is_report_title(line):
        return False
    if re.match(r"^【.+】$", line):
        return True
    if len(line) > HEADING_MAX_LEN:
        return False
    if line.endswith(("。", "；", "，", "！", "？", "）", ")", ":", "：")) and len(line) > 15:
        return False
    if re.match(r"^\d{1,2}月\d{1,2}日", line):
        return False
    if re.match(r"^(背景|策略|进展|下周计划|总结|summary)", line, re.I):
        return False
    if len(line) <= HEADING_MAX_LEN:
        if next_line is None:
            return len(line) <= 20
        # 连续两个短标题（如 catchii房间需求 -> 家族房优先需求）
        if len(next_line) <= HEADING_MAX_LEN and not next_line.endswith("。"):
            return True
        if len(next_line) > len(line) + 10 or next_line.endswith("。"):
            return True
    return False


def get_chunk(
    text: str,
    max_token_len: int = MAX_TOKEN_LEN,
    cover_content: int = COVER_CONTENT,
) -> List[str]:
    chunk_text = []
    curr_len = 0
    curr_chunk = ""
    token_len = max_token_len - cover_content
    lines = text.splitlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
        line_len = len(enc.encode(line))

        if line_len > max_token_len:
            if curr_chunk:
                chunk_text.append(curr_chunk)
                curr_chunk = ""
                curr_len = 0

            line_tokens = enc.encode(line)
            num_chunks = (len(line_tokens) + token_len - 1) // token_len

            for i in range(num_chunks):
                start_token = i * token_len
                end_token = min(start_token + token_len, len(line_tokens))
                chunk_part = enc.decode(line_tokens[start_token:end_token])

                if i > 0:
                    overlap_start = max(0, start_token - cover_content)
                    cover_part = enc.decode(line_tokens[overlap_start:start_token])
                    chunk_part = cover_part + chunk_part

                chunk_text.append(chunk_part)

            curr_chunk = ""
            curr_len = 0

        elif curr_len + line_len + 1 <= token_len:
            if curr_chunk:
                curr_chunk += "\n"
                curr_len += 1
            curr_chunk += line
            curr_len += line_len
        else:
            if curr_chunk:
                chunk_text.append(curr_chunk)
            if chunk_text:
                cover_part = get_token_cover(chunk_text[-1], cover_content)
                curr_chunk = cover_part + "\n" + line
                curr_len = len(enc.encode(cover_part)) + 1 + line_len
            else:
                curr_chunk = line
                curr_len = line_len

    if curr_chunk:
        chunk_text.append(curr_chunk)

    return chunk_text


def split_docx_into_sections(file_path: str) -> List[Tuple[str, List[str]]]:
    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                paragraphs.append(row_text)

    sections: List[Tuple[str, List[str]]] = []
    current_heading = "综合"
    current_lines: List[str] = []

    for i, line in enumerate(paragraphs):
        if is_report_title(line):
            continue
        next_line = paragraphs[i + 1] if i + 1 < len(paragraphs) else None
        if is_section_heading(line, next_line):
            if current_lines or current_heading != "综合":
                sections.append((current_heading, current_lines))
            current_heading = line
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_heading, current_lines))

    if not sections and paragraphs:
        sections.append(("综合", paragraphs))

    return sections


def build_chunk_text(
    source: str,
    report_date: str,
    project: str,
    body_lines: List[str],
) -> str:
    header = (
        f"[来源: {source} | 日期: {format_report_date(report_date)} | 项目: {project}]"
    )
    body = "\n".join([project] + body_lines) if body_lines else project
    return f"{header}\n{body}"


class ReadFiles:
    def __init__(self, path: str) -> None:
        self._path = os.path.abspath(path)
        self.file_list = self.get_files()

    def get_files(self) -> List[str]:
        file_list = []
        for filepath, _, filenames in os.walk(self._path):
            for filename in filenames:
                if filename.startswith("~$") or filename == "tmp.docx":
                    continue
                if filename.endswith(".docx"):
                    file_list.append(os.path.join(filepath, filename))
        return sorted(file_list)

    def get_chunks(self) -> List[DocumentChunk]:
        chunks: List[DocumentChunk] = []
        for file_path in self.file_list:
            rel_path = os.path.relpath(file_path, self._path)
            report_date = parse_report_date_from_path(file_path)
            sections = split_docx_into_sections(file_path)

            if sections and not report_date:
                first_line = sections[0][1][0] if sections[0][1] else sections[0][0]
                report_date = parse_report_date_from_text(first_line)

            for project, body_lines in sections:
                if not body_lines:
                    continue
                section_text = build_chunk_text(rel_path, report_date, project, body_lines)
                token_count = len(enc.encode(section_text))

                if token_count <= MAX_TOKEN_LEN:
                    chunks.append(
                        DocumentChunk(section_text, rel_path, report_date, project)
                    )
                else:
                    for part in get_chunk(section_text):
                        chunks.append(
                            DocumentChunk(part, rel_path, report_date, project)
                        )

        return chunks
